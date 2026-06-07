#!/usr/bin/env python3
"""
将公众号 markdown 文章转为百家号上传格式（DOCX / PDF / TXT）
保留最大程度的结构样式：标题层级、加粗、表格、代码块、列表

用法:
  python3 scripts/convert_to_baijiahao.py 创作/文章_xxx.md
  python3 scripts/convert_to_baijiahao.py 创作/文章_xxx.md --format docx
  python3 scripts/convert_to_baijiahao.py 创作/文章_xxx.md --format all
  python3 scripts/convert_to_baijiahao.py 创作/文章_xxx.md --adapt   # 自动适配为百家号风格
"""
import re
import sys
import json
import argparse
from pathlib import Path

VENV_PYTHON = "/tmp/docx_venv/bin/python3"

# ──────────────────────────────
# TXT 输出
# ──────────────────────────────
def to_txt(md_text: str) -> str:
    """Markdown → 纯文本，保留换行和段落结构"""
    text = md_text
    # 去掉 frontmatter
    text = re.sub(r'^---\n.*?\n---\n', '', text, flags=re.DOTALL)
    # 去掉 markdown 链接标记，保留文字
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # 去掉图片标记
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    # 去掉加粗/斜体标记
    text = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', text)
    # 去掉代码块标记
    text = re.sub(r'```\w*\n', '', text)
    text = text.replace('```', '')
    # 去掉行内代码
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # 去掉表格式分隔线
    text = re.sub(r'^\|:?---?:?\|:?---?:?.*\|$', '', text, flags=re.M)
    # 去掉 html 标签
    text = re.sub(r'<[^>]+>', '', text)
    # 去掉分隔线
    text = re.sub(r'^---+$', '', text, flags=re.M)
    # 去掉标题标记
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.M)
    # 去掉无序列表标记
    text = re.sub(r'^[\s]*[-*+]\s+', '', text, flags=re.M)
    # 去掉有序列表数字
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.M)
    # 清理多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ──────────────────────────────
# DOCX 输出（使用 python-docx）
# ──────────────────────────────
def to_docx(md_text: str, output_path: str, adapt_for_baijiaha: bool = False):
    """
    Markdown → DOCX 转换
    通过子进程调用 venv 中的 python-docx，因为系统 python 环境受限
    """
    import subprocess
    # 将 markdown 文本序列化为 JSON 传给子进程
    payload = json.dumps({
        "md": md_text,
        "output": output_path,
        "adapt": adapt_for_baijiaha,
    })
    script = r'''
import sys, json, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

data = json.loads(sys.argv[1])
md_text = data["md"]
output_path = data["output"]
adapt = data["adapt"]

# 如果要求适配百家号风格，做预处理
if adapt:
    # 把「说句得罪人的话/扎心的」等刺式话术去掉或弱化
    md_text = re.sub(r'说句(得罪人|扎心|不好听)的[话，。!！]?', '', md_text)
    # 把「你选哪个？评论区说说」等互动引导去掉
    md_text = re.sub(r'你选(哪个|A|B|C)[^。\n]*[。\n]', '', md_text)
    md_text = re.sub(r'评论区.*[。\n]', '', md_text)
    # 把「下期预告」「点个关注」等去掉
    md_text = re.sub(r'(还没关注|点个关注|下期预告|下期见|不迷路|别走丢).*', '', md_text)
    # 清理多余空行
    md_text = re.sub(r'\n{3,}', '\n\n', md_text)

doc = Document()

# ── 设置默认字体 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ── 辅助函数 ──
def add_code_block(doc, code_text):
    """添加代码块（灰色背景等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    # 添加灰色底纹
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F0F0')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_table_from_markdown(doc, rows_data):
    """从 markdown 表格行数据创建表格"""
    if len(rows_data) < 2:
        return
    header = [c.strip() for c in rows_data[0].split('|')[1:-1]]
    body = []
    for row in rows_data[2:]:  # 跳过分割线行
        cells = [c.strip() for c in row.split('|')[1:-1]]
        if cells:
            body.append(cells)
    if not body:
        return
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # 数据行
    for ri, row in enumerate(body):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

# ── 解析 markdown 行并构建 DOCX ──
lines = md_text.split('\n')
i = 0
in_code_block = False
code_buffer = []
in_table = False
table_buffer = []

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # 跳过 frontmatter
    if i == 0 and stripped == '---':
        i += 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1
        continue

    # 代码块开始/结束
    if stripped.startswith('```'):
        if in_code_block:
            add_code_block(doc, '\n'.join(code_buffer))
            code_buffer = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue
    if in_code_block:
        code_buffer.append(line)
        i += 1
        continue

    # 表格（连续的行）
    if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
        table_buffer.append(stripped)
        i += 1
        continue
    else:
        if len(table_buffer) >= 2:
            add_table_from_markdown(doc, table_buffer)
            table_buffer = []
        if table_buffer:
            table_buffer = []

    # 空行
    if not stripped:
        i += 1
        continue

    # 分隔线
    if re.match(r'^-{3,}$', stripped):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('─' * 40)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)
        i += 1
        continue

    # 标题
    heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2)
        # 去掉加粗标记
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        heading = doc.add_heading(text, level=level)
        i += 1
        continue

    # 无序列表
    list_match = re.match(r'^[\s]*[-*+]\s+(.*)', stripped)
    if list_match:
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', list_match.group(1))
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.size = Pt(11)
        i += 1
        continue

    # 有序列表（① ② ③ 或 1. 2. 3. 或 数字. ）
    ordered_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
    circled_match = re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]\s+(.*)', stripped)
    if ordered_match:
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', ordered_match.group(2))
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(text)
        run.font.size = Pt(11)
        i += 1
        continue
    if circled_match:
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', circled_match.group(1))
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        run.font.size = Pt(11)
        i += 1
        continue

    # 普通段落（处理加粗、行内代码、链接）
    p = doc.add_paragraph()
    # 用正则解析内联格式
    text = stripped
    # 收集片段
    segments = []
    pos = 0
    for m in re.finditer(r'(\*\*(.+?)\*\*|`([^`]+)`|\[([^\]]+)\]\([^)]+\))', text):
        if m.start() > pos:
            segments.append(('text', text[pos:m.start()]))
        if m.group(2):  # **bold**
            segments.append(('bold', m.group(2)))
        elif m.group(3):  # `code`
            segments.append(('code', m.group(3)))
        elif m.group(4):  # [text](url)
            segments.append(('text', m.group(4)))
        pos = m.end()
    if pos < len(text):
        segments.append(('text', text[pos:]))

    if not segments:
        segments = [('text', text)]

    for seg_type, seg_text in segments:
        run = p.add_run(seg_text)
        run.font.size = Pt(11)
        if seg_type == 'bold':
            run.font.bold = True
        elif seg_type == 'code':
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xE8, 0x6C, 0x00)

    i += 1

# 处理缓冲区中的表格（文件结尾的情况）
if len(table_buffer) >= 2:
    add_table_from_markdown(doc, table_buffer)

doc.save(output_path)
print(f"✅ DOCX saved: {output_path}")
'''
    result = subprocess.run(
        [VENV_PYTHON, '-c', script, payload],
        capture_output=True, text=True, timeout=30
    )
    if result.returncode != 0:
        print(f"❌ DOCX failed: {result.stderr}")
        return False
    if result.stdout:
        print(result.stdout.strip())
    return True


# ──────────────────────────────
# PDF 输出（尝试多种方案）
# ──────────────────────────────
def to_pdf(md_text: str, output_path: str):
    """Markdown → PDF（通过 HTML 中转）"""
    import subprocess, tempfile

    # 先用 markdown 库转 HTML
    try:
        import markdown
        html_body = markdown.markdown(md_text, extensions=['extra'])
    except ImportError:
        print("⚠️ markdown 库未装，用简单转换")
        html_body = f"<pre>{md_text}</pre>"

    # 包装成完整 HTML
    html_full = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>
body {{ font-family: 'Microsoft YaHei', sans-serif; padding: 40px; line-height: 1.8; }}
h1 {{ font-size: 22px; color: #111; }}
h2 {{ font-size: 18px; color: #333; }}
h3 {{ font-size: 16px; color: #555; }}
pre {{ background: #f5f5f5; padding: 16px; border-radius: 6px; font-size: 13px; }}
code {{ background: #f0f0f0; padding: 2px 6px; border-radius: 3px; font-size: 13px; }}
table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
th, td {{ border: 1px solid #ddd; padding: 8px; }}
th {{ background: #f0f0f0; }}
blockquote {{ border-left: 4px solid #ddd; padding-left: 16px; color: #666; }}
</style></head><body>{html_body}</body></html>"""

    # 方案1: wkhtmltopdf
    wkhtml = subprocess.run(['which', 'wkhtmltopdf'], capture_output=True, text=True)
    if wkhtml.returncode == 0:
        with tempfile.NamedTemporaryFile(suffix='.html', mode='w', delete=False, encoding='utf-8') as f:
            f.write(html_full)
            html_path = f.name
        result = subprocess.run(
            ['wkhtmltopdf', '--encoding', 'utf-8', html_path, output_path],
            capture_output=True, text=True, timeout=30
        )
        Path(html_path).unlink(missing_ok=True)
        if result.returncode == 0:
            print(f"✅ PDF saved: {output_path}")
            return True
        else:
            print(f"⚠️ wkhtmltopdf failed: {result.stderr[:200]}")

    # 方案2: weasyprint
    weasy = subprocess.run(['which', 'weasyprint'], capture_output=True, text=True)
    weasy_venv = subprocess.run(
        ['which', 'weasyprint'], capture_output=True, text=True,
        env={**__import__('os').environ, 'PATH': f"/tmp/docx_venv/bin:{__import__('os').environ.get('PATH', '')}"}
    )
    weasy_path = None
    if weasy.returncode == 0:
        weasy_path = 'weasyprint'
    elif weasy_venv.returncode == 0:
        weasy_path = weasy_venv.stdout.strip()
    if weasy_path:
        result = subprocess.run(
            [weasy_path, '-', output_path],
            input=html_full, capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            print(f"✅ PDF saved: {output_path}")
            return True
        else:
            print(f"⚠️ weasyprint failed: {result.stderr[:200]}")

    print("⚠️ PDF 生成失败，请安装 wkhtmltopdf 或 weasyprint")
    print("   apt install wkhtmltopdf   # 或")
    print("   pip install weasyprint")
    return False


# ──────────────────────────────
# 主入口
# ──────────────────────────────
def convert(md_path: str, fmt: str = "docx", output_dir: str = None, adapt: bool = False):
    md_path = Path(md_path)
    if not md_path.exists():
        print(f"❌ 文件不存在: {md_path}")
        return

    md_text = md_path.read_text(encoding='utf-8')

    # 提取标题做文件名
    title_match = re.search(r'^title:\s*(.+?)\s*$', md_text, re.M)
    title = title_match.group(1).strip() if title_match else md_path.stem
    safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)[:40]

    if output_dir:
        out_dir = Path(output_dir)
    else:
        out_dir = md_path.parent / "百家号输出"
    out_dir.mkdir(parents=True, exist_ok=True)

    results = []
    if fmt in ("docx", "all"):
        out = out_dir / f"{safe_title}.docx"
        ok = to_docx(md_text, str(out), adapt)
        results.append(("DOCX", str(out), ok))
    if fmt in ("txt", "all"):
        out = out_dir / f"{safe_title}.txt"
        txt = to_txt(md_text)
        out.write_text(txt, encoding='utf-8')
        print(f"✅ TXT saved: {out}")
        results.append(("TXT", str(out), True))
    if fmt in ("pdf", "all"):
        out = out_dir / f"{safe_title}.pdf"
        ok = to_pdf(md_text, str(out))
        results.append(("PDF", str(out), ok))

    print(f"\n{'='*50}")
    print(f"📄 源文件: {md_path.name}")
    print(f"🏷️  标题: {title}")
    for name, path, ok in results:
        icon = "✅" if ok else "⚠️"
        print(f"  {icon} {name}: {path}")
    print(f"{'='*50}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Markdown → 百家号上传格式")
    parser.add_argument("file", help="文章 markdown 文件路径")
    parser.add_argument("--format", "-f", choices=["docx", "txt", "pdf", "all"],
                        default="docx", help="输出格式（默认docx）")
    parser.add_argument("--output", "-o", help="输出目录（默认 创作/百家号输出/）")
    parser.add_argument("--adapt", "-a", action="store_true",
                        help="适配百家号风格（去刺式话术、去互动引导）")
    args = parser.parse_args()
    convert(args.file, args.format, args.output, args.adapt)
