#!/usr/bin/env python3
"""将公众号文章 Markdown 转为 .docx，存入值得顶目录"""
import sys, re, os
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ 请先安装: pip install python-docx")
    sys.exit(1)

def markdown_to_docx(md_path, output_dir):
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()

    # 提取标题
    title_match = re.search(r'^title:\s*(.+?)\s*$', text, re.M)
    title = title_match.group(1).strip() if title_match else Path(md_path).stem

    # 去除 frontmatter
    body = re.sub(r'^---\n.*?\n---\n', '', text, flags=re.DOTALL)

    doc = Document()

    # 标题
    doc.add_heading(title, level=1)

    # 时间
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f'发布于: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # spacing

    # 正文
    for line in body.split('\n'):
        line = line.strip()
        if not line:
            continue

        # H2
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        # H3
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        # 引用块
        elif line.startswith('> '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:].strip())
            run.font.italic = True
            run.font.color.rgb = RGBColor(107, 114, 128)
        # 分隔线
        elif line.startswith('---') and len(line) <= 4:
            doc.add_paragraph('─' * 40)
        # 列表（无序）
        elif line.startswith('- '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        # 列表（有序）——尝试匹配数字开头
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
        # 普通段落
        else:
            # 处理加粗
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.+?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # 保存
    os.makedirs(output_dir, exist_ok=True)
    safe_title = re.sub(r'[\\/:*?"<>|]', '', title)[:60]
    filename = f'{safe_title}.docx'
    output_path = Path(output_dir) / filename
    doc.save(str(output_path))
    print(f'✅ {filename} ({os.path.getsize(output_path)/1024:.1f} KB)')
    return output_path


if __name__ == '__main__':
    md_path = sys.argv[1] if len(sys.argv) > 1 else None
    output_dir = sys.argv[2] if len(sys.argv) > 2 else '值得顶'

    if not md_path:
        # 自动找最新的文章
        创作_dir = Path(__file__).parent / '创作'
        files = sorted(创作_dir.glob('文章_*.md'), key=lambda p: p.stat().st_mtime, reverse=True)
        md_path = str(files[0]) if files else None

    if md_path and Path(md_path).exists():
        markdown_to_docx(md_path, output_dir)
    else:
        print('❌ 未找到文章文件')
